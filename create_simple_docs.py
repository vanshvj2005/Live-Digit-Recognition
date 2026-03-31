from docx import Document

def create_document():
    doc = Document()
    
    # Title
    doc.add_heading('Live Webcam Handwritten Digit Recognition', 0)
    
    # 1. Intro
    doc.add_heading('1. Project Introduction', level=1)
    doc.add_paragraph(
        "The objective of this project is to develop a real-time computer vision system capable "
        "of identifying handwritten digits (0-9). By leveraging a live webcam feed, the system processes "
        "physical handwriting on paper, isolates the character from background noise, and uses a trained "
        "Deep Learning (CNN) model to accurately classify the digit."
    )
    
    # 2. Technologies
    doc.add_heading('2. Core Technologies', level=1)
    doc.add_paragraph(
        "• OpenCV: Used for real-time video capture, image processing, cropping, and contour detection.\n"
        "• TensorFlow/Keras: Runs the Convolutional Neural Network (CNN) model to predict digits.\n"
        "• NumPy: Performs matrix calculations and thresholding."
    )
    
    # 3. Features
    doc.add_heading('3. Key Features & Excel Data Logging', level=1)
    doc.add_paragraph(
        "One of the standout features of this project is its robust filtering and real-time logging. To track the "
        "system's performance, the software logs live prediction data so it can be viewed in Excel. The logged data includes:\n"
        "• Timestamp: The exact time the digit was processed.\n"
        "• Predicted Digit: The neural network's final output.\n"
        "• Confidence Score (%): The probability accuracy of the prediction.\n"
        "• Fill Ratio: How much ink is present in the bounding box, used to eliminate false positives like faces."
    )
    
    # 4. Conclusion
    doc.add_heading('4. Conclusion', level=1)
    doc.add_paragraph(
        "This project successfully brings real-world objects into the digital space, combining reliable "
        "machine learning models with strong computer vision foundations to prevent errors."
    )
    
    # Save the file
    file_name = 'Simple_Project_Report.docx'
    doc.save(file_name)
    print(f"Successfully generated {file_name}")

if __name__ == '__main__':
    create_document()
