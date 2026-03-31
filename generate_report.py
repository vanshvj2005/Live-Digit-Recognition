from docx import Document
from docx.shared import Pt, Inches

def create_document():
    doc = Document()
    
    # Title
    doc.add_heading('Live Digit Recognition - Updates and Fixes', 0)
    
    # Intro
    doc.add_paragraph(
        "This document outlines the recent fixes made to the live_digit_recognition.py script to drastically "
        "improve the accuracy and reliability of real-time handwritten digit recognition via a webcam feed. "
        "The goal was to prevent false positives (like analyzing the user's face or the room background) and "
        "properly read very faint, thin pen strokes drawn on lined paper."
    )
    
    # Section: What I Did
    doc.add_heading('Key System Updates', level=1)
    
    doc.add_heading('1. Central Scan Region (ROI Box)', level=2)
    doc.add_paragraph(
        "The program used to scan the entire 640px wide webcam feed for anything that slightly resembled "
        "a solid line. Now, it generates a fixed 300x300 bounding box right in the center of the screen ("
        "'Place digit here'). We mathematically crop the video feed before thresholding so that all background "
        "objects in the user's room, including shadows on the walls, are physically excluded from the Neural Network's view."
    )
    
    doc.add_heading('2. Advanced Mathematical "Fill Ratio" Validation', level=2)
    doc.add_paragraph(
        "Faces have extremely soft shadows and complex features that produce thick white noise blobs under "
        "OpenCV Adaptive Thresholding. To combat this, I added a 'Fill Ratio' variable. The program now calculates "
        "the exact percentage of white 'ink' pixels versus empty black space inside the bounding box. "
        "A drawn digit usually takes up roughly 2% to 65% of the box, whereas faces take up 80%+ and stray dots "
        "take less than 1%. If the shape fails the mathematical fill ratio test, it is ignored."
    )
    
    doc.add_heading('3. Minimum Width & Aspect Ratio Relaxations', level=2)
    doc.add_paragraph(
        "The algorithm was originally trained for thick, square numbers like '8' and '3'. A perfectly straight "
        "line like the number '1' drawn by a pencil is incredibly tall and very thin. It was being ignored "
        "because its bounding box was less than 15 pixels wide! I relaxed the required aspect ratio "
        "constraints, allowing widths as small as 5 pixels, so vertical lines are perfectly detected."
    )
    
    doc.add_heading('4. Morphological 'Dilation' for Faint Ink', level=2)
    doc.add_paragraph(
        "The CNN was trained on the MNIST Dataset, which consists of numbers drawn with the equivalent thickness "
        "of large markers. When someone writes a massive number using an ultra-fine pen, the neural network "
        "is confused by the extremely microscopic width of the line. Before sending the threshold mask to the "
        "Neural Network, I used `cv2.dilate` with a 5x5 mathematical kernel. This artificially takes the thin pen "
        "stroke and thickens the edges uniformly in all directions. The faint 1-pixel pen string is converted "
        "into a thick, highly-confident 5-pixel stroke without destroying the shape."
    )

    # Section: Libraries Used
    doc.add_heading('Libraries & Technologies Used', level=1)
    
    libraries = [
        ("cv2 (OpenCV)", "Used extensively for the Computer Vision pipeline. Key functions included checking the webcam feed (cv2.VideoCapture), rendering visual guides (cv2.rectangle), converting colors to Grayscale (cv2.cvtColor), blurring out artifacts (cv2.GaussianBlur), applying inverted mathematical thresholds (cv2.adaptiveThreshold), thickening lines (cv2.dilate), and isolating localized shapes (cv2.findContours & cv2.boundingRect)."),
        ("numpy (np)", "Used mainly for highly efficient matrix and array operations, which are the backbone of all computer vision code. Here, numpy specifically defines the 5x5 matrix kernel shape used during cv2.dilate to thicken lines. It is also used (np.argmax, np.max) to evaluate the output probability array coming from the CNN model and retrieve the highest-predicted class and confidence level."),
        ("tensorflow.keras", "The framework that holds the trained Artificial Intelligence. `tensorflow.keras.models.load_model()` safely injects the previous HDF5 (`.h5`) deep learning checkpoint into our live script, and `model.predict()` executes the lightning fast forward-pass calculation across the GPU/CPU to classify the 28x28 preprocessed box."),
        ("imutils", "A heavily utilized wrapper library for basic image processing tasks. Specifically deployed to effortlessly resize the webcam's native dimension footprint instantly down to `width=640` while maintaining the exact aspect ratio mathematically without explicit calculations block down the road.")
    ]
    
    for lib, desc in libraries:
        p = doc.add_paragraph()
        p.add_run(f"• {lib}:").bold = True
        p.add_run(f" {desc}")
    
    doc.add_paragraph(
        "\nAll updates have successfully synchronized the real-world properties of faint script with "
        "the specific requirements of the underlying CNN model."
    )
    
    file_name = 'Digit_Recognition_Updates_Report.docx'
    doc.save(file_name)
    print(f"Successfully generated {file_name}")

if __name__ == '__main__':
    create_document()
